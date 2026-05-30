import csv
import io
import logging
from pathlib import Path

import fitz  # PyMuPDF
import openpyxl
from docx import Document as DocxDocument
from langchain_community.document_loaders import (
    PDFPlumberLoader,
    PyMuPDFLoader,
    TextLoader,
)
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".pdf", ".txt", ".docx", ".csv", ".xlsx", ".xls", ".doc",
}


class DocumentLoaderService:
    """统一文档加载器。

    支持格式：
      PDF  — 三级回退策略 (PyMuPDF → PDFPlumber → fitz 逐页)
      TXT  — TextLoader (UTF-8)
      DOCX — python-docx 直接提取段落文字
      DOC  — 由 olefile 提取 OLE 复合文档中的文字流
      CSV  — Python csv 模块逐行读取
      XLSX — openpyxl 逐单元格提取
      XLS  — openpyxl (仅支持新版 .xls，旧版 97-2003 需 xlrd)
    """

    @staticmethod
    def load(file_path: str | Path) -> list[Document]:
        path = Path(file_path)
        if not path.exists():
            raise ValueError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"不支持的文件格式 '{ext}'，"
                f"仅支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        logger.info("加载文档: %s (格式=%s)", path.name, ext)

        # 路由到对应 loader
        if ext == ".pdf":
            docs = _load_pdf(str(path))
        elif ext == ".txt":
            docs = _load_txt(str(path))
        elif ext == ".docx":
            docs = _load_docx(str(path))
        elif ext == ".doc":
            docs = _load_doc(str(path))
        elif ext == ".csv":
            docs = _load_csv(str(path))
        elif ext in (".xlsx", ".xls"):
            docs = _load_xlsx(str(path))
        else:
            docs = []

        # 过滤空页面
        docs = [d for d in docs if d.page_content and d.page_content.strip()]
        if not docs:
            logger.warning(
                "文档 %s 未提取到任何有效文字", path.name,
            )
            return []

        for doc in docs:
            doc.metadata["filename"] = path.name

        logger.info("文档加载完成: %s, 共 %d 片段", path.name, len(docs))
        return docs


# ── TXT ────────────────────────────────────────────────────

def _load_txt(file_path: str) -> list[Document]:
    try:
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()
    except Exception:
        logger.exception("TXT 解析失败: %s", file_path)
        raise RuntimeError(f"TXT 文件解析失败: {file_path}") from None


# ── PDF ────────────────────────────────────────────────────

def _load_pdf(file_path: str) -> list[Document]:
    try:
        loader = PyMuPDFLoader(file_path)
        docs = loader.load()
        logger.info("PDF 由 PyMuPDFLoader 成功加载")
        return docs
    except Exception:
        logger.warning("PyMuPDFLoader 失败，切换到 PDFPlumberLoader ...")

    try:
        loader = PDFPlumberLoader(file_path)
        docs = loader.load()
        logger.info("PDF 由 PDFPlumberLoader 成功加载")
        return docs
    except Exception:
        logger.warning("PDFPlumberLoader 也失败，切换到 fitz 逐页容错 ...")

    try:
        docs = _load_pdf_page_by_page(file_path)
        if docs:
            logger.info("PDF 由 fitz 逐页容错成功加载 (%d 页)", len(docs))
            return docs
    except Exception:
        pass

    raise RuntimeError(
        f"PDF 解析失败: {Path(file_path).name}。"
        f"已尝试 PyMuPDF / PDFPlumber / fitz 三种方式均无法解析。"
    )


def _load_pdf_page_by_page(file_path: str) -> list[Document]:
    docs: list[Document] = []
    failed_pages: list[int] = []
    try:
        pdf = fitz.open(file_path)
    except Exception:
        raise RuntimeError(f"fitz 无法打开 PDF 文件: {file_path}")

    for page_num in range(len(pdf)):
        try:
            page = pdf[page_num]
            text = page.get_text()
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"page": page_num, "source": file_path},
                ))
        except Exception:
            failed_pages.append(page_num + 1)

    pdf.close()

    if failed_pages:
        logger.warning("fitz 逐页提取时 %d 页失败，已跳过", len(failed_pages))

    return docs


# ── DOCX ───────────────────────────────────────────────────

def _load_docx(file_path: str) -> list[Document]:
    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            logger.warning("DOCX 文件无文字: %s", file_path)
            return []
        text = "\n\n".join(paragraphs)
        return [Document(page_content=text, metadata={"source": file_path})]
    except Exception:
        logger.exception("DOCX 解析失败: %s", file_path)
        raise RuntimeError(f"DOCX 文件解析失败: {Path(file_path).name}") from None


# ── DOC (旧格式) ────────────────────────────────────────────

def _load_doc(file_path: str) -> list[Document]:
    """解析旧版 .doc 文件。

    策略：
      1. 先尝试作为 .docx 打开（很多 .doc 实际是新格式）
      2. 用 olefile 读取 OLE 容器中的 WordDocument 主文字流
    """
    import olefile

    # 级别 1：尝试作为 docx 打开（覆盖大多数 .doc 文件）
    try:
        return _load_docx(file_path)
    except Exception:
        pass

    # 级别 2：OLE 复合文档 — 只读文字流
    try:
        if not olefile.isOleFile(file_path):
            raise RuntimeError("不是有效的 OLE 复合文档")
    except Exception as exc:
        raise RuntimeError(
            f"DOC 文件解析失败: {Path(file_path).name}。"
            f"文件不是有效的 DOC 或 DOCX 格式。"
        ) from exc

    try:
        ole = olefile.OleFileIO(file_path)
        text_parts: list[str] = []

        # 只读取已知包含文字的流（避免遍历全部二进制资源）
        text_streams = [
            "WordDocument",           # 正文
            "1Table", "0Table",       # 辅助文字表
        ]
        for name in text_streams:
            try:
                if ole.exists(name):
                    data = ole.openstream(name).read()
                    # Word 文档内部编码通常是单字节 ASCII 扩展或 UTF-16LE
                    text = _decode_word_binary(data)
                    if text:
                        text_parts.append(text)
            except Exception:
                continue

        # 如果主要流没提取到文字，尝试读 metadata
        if not text_parts:
            meta_text = _extract_ole_metadata(ole)
            if meta_text:
                text_parts.append(meta_text)

        ole.close()

        if not text_parts:
            raise RuntimeError("无法从 DOC 文件中提取任何文字。")

        full_text = "\n\n".join(text_parts)
        return [Document(page_content=full_text, metadata={"source": file_path})]

    except RuntimeError:
        raise
    except Exception:
        logger.exception("DOC 解析失败: %s", file_path)
        raise RuntimeError(f"DOC 文件解析失败: {Path(file_path).name}") from None


def _decode_word_binary(data: bytes, max_len: int = 1_000_000) -> str:
    """从 Word OLE 二进制数据中提取可打印文字。

    尝试 UTF-16LE 解码（Unicode Word），失败则用 Latin-1（ASCII 扩展）。
    限制大小防止处理巨型嵌入图片/字体。
    """
    if len(data) > max_len:
        data = data[:max_len]
    # 尝试 Unicode 双字节编码
    try:
        text = data.decode("utf-16-le", errors="replace")
    except Exception:
        text = data.decode("latin-1", errors="replace")
    # 过滤：只保留含连续可打印字符的片段
    return _filter_printable(text, min_len=40)


def _extract_ole_metadata(ole) -> str:
    """从 OLE 元数据流提取标题/主题等文字。"""
    parts = []
    for prop in ["\x05SummaryInformation", "\x05DocumentSummaryInformation"]:
        try:
            if ole.exists(prop):
                data = ole.openstream(prop).read()
                text = data.decode("latin-1", errors="replace")
                printable = "".join(c for c in text if c.isprintable() or c in "\n\r\t ")
                if len(printable) > 20:
                    parts.append(printable)
        except Exception:
            continue
    return "\n".join(parts)


def _filter_printable(text: str, min_len: int = 40) -> str:
    """从文本中提取连续可打印字符片段，过滤二进制噪声。"""
    result: list[str] = []
    buf: list[str] = []
    for c in text:
        if c.isprintable() or c in "\n\r\t ":
            buf.append(c)
        else:
            if len(buf) >= min_len:
                result.append("".join(buf))
            buf = []
    if len(buf) >= min_len:
        result.append("".join(buf))
    return "\n".join(result)


# ── CSV ────────────────────────────────────────────────────

def _load_csv(file_path: str) -> list[Document]:
    try:
        # 先尝试 UTF-8，失败则用 GBK
        try:
            with open(file_path, encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, encoding="gbk") as f:
                content = f.read()

        reader = csv.reader(io.StringIO(content))
        rows = list(reader)
        if not rows:
            return []

        # 每行用 tab 连接各列，行与行之间换行
        lines = [" | ".join(r) for r in rows if any(c.strip() for c in r)]
        text = "\n".join(lines)
        return [Document(page_content=text, metadata={"source": file_path})]
    except Exception:
        logger.exception("CSV 解析失败: %s", file_path)
        raise RuntimeError(f"CSV 文件解析失败: {Path(file_path).name}") from None


# ── XLSX / XLS ─────────────────────────────────────────────

def _load_xlsx(file_path: str) -> list[Document]:
    sheets_text: list[str] = []
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    except Exception:
        logger.exception("openpyxl 无法打开: %s", file_path)
        raise RuntimeError(f"Excel 文件解析失败: {Path(file_path).name}") from None

    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append(" | ".join(cells))
                    row_count += 1
            if rows:
                sheet_header = f"--- 工作表: {sheet_name} ---"
                sheets_text.append(sheet_header + "\n" + "\n".join(rows))
            logger.info("工作表 '%s': %d 行", sheet_name, row_count)
    finally:
        wb.close()

    if not sheets_text:
        return []
    full_text = "\n\n".join(sheets_text)
    return [Document(page_content=full_text, metadata={"source": file_path})]
